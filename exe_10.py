#CRUD#
import msvcrt
from docx import Document
import os


usuarios = [];
id_inicial = -1

def adicionar_Usuario(id, name, idade, cpf):
    usuario = {"id": id, "nome": name, "idade":idade, "cpf":cpf}

    usuarios.append(usuario)

def remover_Usuario(name):
    for i in usuarios:
        if i["nome"] == name:
            usuarios.remove(i)
            break

def consultarUsuarios():
    for i in usuarios:
        print(f'{i}')    

def input_cpf():
    cpf = ''
    formatado = ''
    print('CPF do usuário: ', end='', flush=True)

    while len(cpf) < 11:
        digito = msvcrt.getwch()

        if digito.isdigit():
            cpf += digito
            formatado += digito
            print(digito, end='', flush=True)

            if len(cpf) % 3 == 0 and len(cpf) < 11:
                formatado += '-'
                print('-', end='', flush=True)

    print()
    return formatado

def editar_Usuarios(name):
    for i in usuarios:
        if i["nome"] == name:
            print('\n' + '═' * 40)
            print(f'Escolha abaixo o que deseja editar de(a) {i["nome"]}')
            print('1 - Nome')
            print('2 - Idade')
            print('3 - cpf')

            new_Data = int(input("Numero: "))
            if new_Data == 1:
                new_name = input("Novo nome:")
                i["nome"] = new_name
            if new_Data == 2:
                new_idade = int(input("Nova idade:"))
                i["idade"] = new_idade

            if new_Data == 3:
                new_cpf = input_cpf()
                i["cpf"] = new_cpf

def salvar_Usuarios():
    doc = Document()
    doc.add_heading("Lista de Usuários",0)

    for u in usuarios:
        doc.add_paragraph(f"ID: {u['id']} | Nome: {u['nome']}| idade :{u['idade']} |CPF: {u['cpf']}")

    desktop = os.path.join(os.path.expanduser('~'),"OneDrive", "Desktop", "usuarios.docx")

    doc.save(desktop)

    

#Iniciar o processo
init = input("Deseja inificar o processo ? (S/N)").lower()
control = True
if init  == 's':
    while control:
        print('\n' + '═' * 40)
        print('        🗂  GERENCIADOR DE USUÁRIOS')
        print('═' * 40)
        print('  1 - ➕  Adicionar usuário')
        print('  2 - ❌  Remover usuário')
        print('  3 - 🔍  Consultar usuários')
        print('  4 - ✏️  Editar')
        print('  5 - 🖨️ Salvar em documento')
        print('═' * 40)
        number  = int(input('Escolha:'))

        if number == 1:
            id_inicial+=1
            nome  = input('Nome do usuario:')

            while not nome.replace(' ','').isalpha():
                print('Nome invalido! Use apenas letras')
                nome  = input('Nome do usuario:')

            try:
                while True:
                    idade_User = int(input('Idade do usuário:'))
                    if idade_User > 100:
                        print('A idade deve ser menor que 100')
                        continue
                    if len(str(idade_User)) > 3:
                        print('A quantidade de digitos deve ser menor que 4')
                        continue
                    else:
                        break

            except ValueError:
                print(' 😢 Erro na idade! Digite apenas números. Tente novamente')
                continue
            
            cpf = input_cpf()

            adicionar_Usuario(id_inicial, nome, idade_User, cpf);

        if number == 2:
            nome = input('Nome do usuário:')
            remover_Usuario(nome);


        if number == 3:
            consultarUsuarios()
        
        
        if number == 4:
            nome = input('Nome do usuario:')
            editar_Usuarios(nome)

        if number == 5:
            salvar_Usuarios()


        #Finalizar o processo 
        finish = input('Deseja finalizar o processo ?').lower()
        if finish == 's':
            control = False





print('\n' + '═' * 40)
print('  ✔  Sistema encerrado com sucesso!')
print(f'  Usuários cadastrados: {len(usuarios)}')
print('  Obrigado por usar o sistema! Até logo 👋')
print('═' * 40 + '\n')